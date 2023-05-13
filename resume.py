from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

from src.Style import *
import src.resume_metadata as data
from src.gui import gui

def Heading(document, heading_content) -> None:
    #First Line
    name = document.add_paragraph(heading_content.name + ", " + heading_content.nickname)
    name_format = name.paragraph_format
    name_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_format.space_before = 0
    name_format.space_after = Cm(0.2)
    name_format.line_spacing = 1
    name.runs[0].bold = True
    name.runs[0].font.size = Pt(26)
    
    #Second Line
    secondLine = document.add_paragraph("Telephone: ")
    secondLine.add_run(heading_content.phone)
    secondLine.add_run(" Email: ")
    secondLine.add_run(heading_content.email)
    HeadingStyle(secondLine)

    #Third Line
    thirdLine = document.add_paragraph("Github: ")
    thirdLine.add_run(heading_content.github)
    thirdLine.add_run(" Website: ")
    thirdLine.add_run(heading_content.website)
    HeadingStyle(thirdLine)

def Objective(document, objective_content) -> None:
    heading = document.add_paragraph("OBJECTIVE")
    objective = document.add_paragraph(objective_content.objective + "\n")
    objective_format = objective.paragraph_format
    objective_format.space_before = 0
    objective_format.space_after = 0
    objective_format.line_spacing = data.CONTENTLINESPACE
    objective.runs[0].bold = False
    objective.runs[0].font.size = Pt(11)
    SubHeadingStyle(heading)

def Education(document, education_content) -> None:
    # Subheading
    heading = document.add_paragraph("EDUCATION")
    SubHeadingStyle(heading)
    # Add Content
    for component in education_content:
        education = document.add_paragraph(component.title)
        table = document.add_table(rows=1, cols=2)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = component.subTitle
        if component.start and component.end:
            heading_cells[1].text = component.start + " - " + component.end
        else:
            heading_cells[1].text = ""
        TableStyle(table)

        ContentDescriptionStyle(document, component.description)
        ContentHeadingStyle(education)

def SideProjects(document, sideProjects_content) -> None:
  # Subheading
    sideProjects = document.add_paragraph("SIDE PROJECT")
    SubHeadingStyle(sideProjects)
    # Add Content
    for component in sideProjects_content:
        if component.subTitle != None:
            sideProject = document.add_paragraph(component.title)
            table = document.add_table(rows=1, cols=2)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = component.subTitle
            if component.start and component.end:
                heading_cells[1].text = component.start + " - " + component.end
            else:
                heading_cells[1].text = ""
            ContentHeadingStyle(sideProject)
        else:
            table = document.add_table(rows=1, cols=2)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = component.title
            if component.start and component.end:
                heading_cells[1].text = component.start + " - " + component.end
            else:
                heading_cells[1].text = ""
            ContentHeadingStyle(heading_cells[0].paragraphs[0])
        TableStyle(table)
        ContentDescriptionStyle(document, component.description)

def Experience(document, experience_content) -> None:
    # Subheading
    experience = document.add_paragraph("EXPERIENCE")
    SubHeadingStyle(experience)

    # Add Content
    for component in experience_content:
        experience = document.add_paragraph(component.title)
        table = document.add_table(rows=1, cols=2)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = component.subTitle
        if component.start and component.end:
            heading_cells[1].text = component.start + " - " + component.end
        else:
            heading_cells[1].text = ""
        TableStyle(table)
        ContentHeadingStyle(experience)
        ContentDescriptionStyle(document, component.description)
        

def Skills(document, skills_content) -> None:
    # Subheading
    skills = document.add_paragraph("SKILLS")
    SubHeadingStyle(skills)

    # Add Content
    table = document.add_table(rows=0, cols=2)
    for skill in skills_content:
        table.add_row()
        heading_cells = table.rows[-1].cells
        heading_cells[0].text = skill.category
        heading_cells[1].text = (skill.list)
        SkillHeadingStyle(heading_cells[0].paragraphs[0])

    SkillStyle(table)

def main(console_message:str = None) -> None:
    # Call gui
    user_data = gui(console_message)
    if not user_data:
        print("Exited")
        return
    heading_data = user_data[0]
    objective_data = user_data[1]
    education_data = user_data[2]
    side_projects_data = user_data[3]
    experience_data = user_data[4]
    skills_data = user_data[5]

    # Create Document
    document = Document()
    style = document.styles["Normal"]
    style.paragraph_format.space_before = Cm(0.2)
    style.paragraph_format.space_after = 0
    style.paragraph_format.line_spacing = 1.5
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    font.bold = False

    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Mm(7.5)
        section.right_margin = Mm(7.5)
        # A4 size
        section.page_height = Mm(297)
        section.page_width = Mm(210)

    Heading(document, heading_data)
    Objective(document, objective_data)
    Education(document, education_data)
    SideProjects(document, side_projects_data)
    Experience(document, experience_data)
    Skills(document, skills_data)

    # Meta Data
    document.core_properties.author = heading_data.name
    document.core_properties.title = heading_data.name + " Resume"
    document.core_properties.subject = "Resume"
    document.core_properties.keywords = "Resume, CV, " + heading_data.name
    document.core_properties.category = "Resume"
    document.core_properties.language = "en-US"
    
    # Save File
    try:
        document.save(heading_data.name + "_resume.docx")
        shell_process = subprocess.Popen([heading_data.name + "_resume.docx"], shell=True) 
        shell_process.wait()
        print("File saved successfully")
    except:
        res = "Error: Unable to save file, please close the file and try again"
        print(res)
        main(res)

if __name__ == "__main__":
    __author__ = "TSE Hui Tung"
    __copyright__ = "Copyright 2023, Formal Resume Generator"
    __license__ = "MIT"
    __version__ = "1.0.0"
    __maintainer__ = "TSE Hui Tung"
    __email__ = "tung23966373@gmail.com"
    __status__ = "Production"
    main()