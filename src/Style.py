from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, RGBColor

from src.horizontal_line import insertHR
import src.resume_metadata as data

def HeadingStyle(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = data.HEADINGLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(11)
    paragraph.runs[1].bold = False
    paragraph.runs[1].font.size = Pt(11)
    paragraph.runs[2].bold = True
    paragraph.runs[2].font.size = Pt(11)
    paragraph.runs[3].bold = False
    paragraph.runs[3].font.size = Pt(11)

def SubHeadingStyle(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = Cm(0.2)
    paragraph_format.line_spacing = data.SUBHEADINGLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(14)
    insertHR(paragraph)

def ContentHeadingStyle(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = data.CONTENTLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(13)

def TableStyle(table) -> None: 
    table.autofit = False 
    table.allow_autofit = False
    table.cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table0_format = table.cell(0,0).paragraphs[0].paragraph_format
    table1_format = table.cell(0,1).paragraphs[0].paragraph_format
    table0_format.space_before = 0
    table0_format.space_after = 0
    table0_format.line_spacing = data.CONTENTLINESPACE
    table1_format.space_before = 0
    table1_format.space_after = 0
    table1_format.line_spacing = data.CONTENTLINESPACE
    table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(0,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in table.rows:
        row.height = Cm(0.1)

def BulletListStyle(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = data.CONTENTLINESPACE

def ContentDescriptionStyle(document, description_list:list) -> None:
    for description in description_list:
            description = document.add_paragraph(description, style="List Bullet")
            BulletListStyle(description)
        
    document.add_paragraph("")

def SkillStyle(table) -> None:
    table.autofit = False 
    table.allow_autofit = False
    for row in table.rows:
        # _EMUS_PER_CM = 360000
        # Full Width = 7052310
        row.height = Cm(0.1)
        row.cells[0].width = Cm(data.SKILLWIDTH)
        row.cells[1].width = 7052310 - data.SKILLWIDTH * 360000

        col0_format = row.cells[0].paragraphs[0].paragraph_format
        col1_format = row.cells[1].paragraphs[0].paragraph_format
        col0_format.space_before = 0
        col0_format.space_after = 0
        col0_format.line_spacing = data.CONTENTLINESPACE
        col1_format.space_before = 0
        col1_format.space_after = 0
        col1_format.line_spacing = data.CONTENTLINESPACE

def SkillHeadingStyle(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = data.CONTENTLINESPACE
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(11)